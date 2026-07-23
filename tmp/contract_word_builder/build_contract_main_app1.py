from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Cm, Pt


ROOT = Path(r"C:\Users\alvad\Documents\Sveton")
SOURCE = ROOT / "01_docs" / "operations" / "contracts" / "Contract_revier.md"
OUTPUT = ROOT / "01_docs" / "operations" / "contracts" / "Contract_revier_main_appendix_1.docx"


def unescape_markdown(value: str) -> str:
    value = value.replace(r"\_", "_")
    value = value.replace(r"\#", "#")
    value = value.replace(r"\|", "|")
    value = value.replace(r"\\", "\\")
    return value


def add_runs(paragraph, text: str, *, font_size: Pt | None = None, bold: bool = False) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if is_bold else part
        run = paragraph.add_run(unescape_markdown(clean))
        run.bold = bold or is_bold
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        if font_size is not None:
            run.font.size = font_size


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    add_runs(paragraph, text.strip(), font_size=Pt(10), bold=bold)


def is_table_separator(line: str) -> bool:
    stripped = line.strip().strip("|").strip()
    return bool(stripped) and all(ch in "-:| " for ch in stripped)


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def add_table(document: Document, raw_rows: list[str]) -> None:
    rows = [split_table_row(row) for row in raw_rows if not is_table_separator(row)]
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    for row in rows:
        row.extend([""] * (max_cols - len(row)))

    table = document.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            set_cell_text(cell, value, bold=row_index == 0)
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")

    document.add_paragraph()


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def add_signature_block(document: Document) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(8)
    spacer.paragraph_format.space_after = Pt(2)

    title = document.add_paragraph()
    title.paragraph_format.keep_with_next = True
    add_runs(title, "Подписи Сторон:", font_size=Pt(11), bold=True)

    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    remove_table_borders(table)
    values = [
        "Заказчик: __________________ /______________/",
        "Исполнитель: _______________ /______________/",
    ]
    for cell, value in zip(table.rows[0].cells, values):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        add_runs(paragraph, value, font_size=Pt(10))


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1.5)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(4)

    for style_name, size in (("Heading 1", 14), ("Heading 2", 12)):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)

    bullet = styles["List Bullet"]
    bullet.font.name = "Times New Roman"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    bullet.font.size = Pt(11)


def add_heading(document: Document, line: str, level: int) -> None:
    text = line[level + 1 :].strip()
    paragraph = document.add_paragraph(style=f"Heading {min(level, 2)}")
    add_runs(paragraph, text, font_size=Pt(14 if level == 1 else 12), bold=True)

    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(document: Document, line: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(paragraph, line.strip(), font_size=Pt(11))


def add_bullet(document: Document, line: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Cm(0.7)
    paragraph.paragraph_format.first_line_indent = Cm(-0.35)
    add_runs(paragraph, line.strip()[2:].strip(), font_size=Pt(11))


def source_slice() -> list[str]:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(lines) if line.strip() == "# ДОГОВОР N \\__\\_")
    end = next(i for i, line in enumerate(lines) if line.strip() == "# Приложение N 2")
    return lines[start:end]


def build() -> None:
    document = Document()
    configure_document(document)

    lines = source_slice()
    table_buffer: list[str] = []

    for raw_line in lines:
        line = raw_line.rstrip()

        if line.strip().startswith("|"):
            table_buffer.append(line)
            continue

        if table_buffer:
            add_table(document, table_buffer)
            table_buffer = []

        stripped = line.strip()
        if not stripped:
            continue
        if stripped == "---":
            continue
        if stripped == "# Приложение N 1":
            document.add_section(WD_SECTION.NEW_PAGE)
            add_heading(document, stripped, 1)
            continue
        if stripped == "## 15. Реквизиты и подписи":
            document.add_page_break()
            add_heading(document, stripped, 2)
            continue
        if stripped.startswith("## "):
            add_heading(document, stripped, 2)
            continue
        if stripped.startswith("# "):
            add_heading(document, stripped, 1)
            continue
        if stripped.startswith("- "):
            add_bullet(document, stripped)
            continue

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if stripped.startswith("г. ") or stripped.startswith('"') or stripped.startswith("к договору") else WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(paragraph, stripped, font_size=Pt(11))

    if table_buffer:
        add_table(document, table_buffer)

    add_signature_block(document)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
