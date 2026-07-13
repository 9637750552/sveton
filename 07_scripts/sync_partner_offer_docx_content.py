from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "01_docs/legal/Коммерческое_предложение_партнерам_Светон.docx"


ROLE_ROW_OLD = "Проверяет условия подключения, щиты, ввод, место установки, нагрузки, ограничения, делает фото/видео и передает исходные данные для расчета и подбора решения."
ROLE_ROW_NEW = "Проводит технический осмотр объекта, проверяет условия подключения, щиты, ввод, место установки, нагрузки и ограничения, делает фото/видео и фиксирует результаты осмотра."

EXCLUSIVE_RULE = "По одному клиенту Светон выплачивает вознаграждение только одному партнеру. Одновременное применение форматов Тип А и Тип Б к одному клиенту не допускается."
MONTAGE_RULE = "Если по объекту возникает необходимость монтажа, подготовительных или иных работ, такие работы согласуются между клиентом и партнером отдельно и оплачиваются клиентом напрямую партнеру."


def set_run_font(run) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(12.5)
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), "Calibri")


def build_paragraph_element(text: str):
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "140")
    spacing.set(qn("w:line"), "259")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    p.append(p_pr)

    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    r_fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), "Calibri")
    r_pr.append(r_fonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    r_pr.append(color)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "25")
    r_pr.append(size)

    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), "25")
    r_pr.append(size_cs)

    r.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def main() -> None:
    doc = Document(str(DOCX_PATH))

    tables = doc.tables
    if len(tables) < 3:
        raise RuntimeError("Expected at least 3 tables in the document.")

    roles_table = tables[1]
    formats_table = tables[2]

    # Update the technical inspection row wording in the roles table.
    for row in roles_table.rows[1:]:
        if row.cells[0].text.strip() == "Технический осмотр и сбор исходных данных":
            row.cells[1].text = ROLE_ROW_NEW
            row.cells[2].text = "Светон выплачивает вознаграждение за технический осмотр объекта и отчет по результатам осмотра."
            for cell in (row.cells[1], row.cells[2]):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run)
            break
    else:
        raise RuntimeError("Could not find the technical inspection row in the roles table.")

    # Ensure the exclusivity rule and montage note sit immediately after the formats table.
    body = doc._body._body
    paragraphs_to_remove = []
    for p in list(body):
        if p.tag != qn("w:p"):
            continue
        text = "".join(p.itertext()).strip()
        if text in {EXCLUSIVE_RULE, MONTAGE_RULE}:
            paragraphs_to_remove.append(p)
    for p in paragraphs_to_remove:
        body.remove(p)

    exclusive_paragraph = build_paragraph_element(EXCLUSIVE_RULE)
    montage_paragraph = build_paragraph_element(MONTAGE_RULE)
    formats_table._tbl.addnext(exclusive_paragraph)
    exclusive_paragraph.addnext(montage_paragraph)

    doc.save(str(DOCX_PATH))


if __name__ == "__main__":
    main()
