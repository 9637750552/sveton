from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "01_docs/legal/Коммерческое_предложение_партнерам_Светон.docx"
LOGO_PATH = ROOT / "01_docs/legal/assets/Logo_Sveton_IBP_2_white.png"


def set_run_font(run, *, size: float, color: RGBColor, bold: bool = False) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), "Calibri")


def clear_container(paragraphs, tables) -> None:
    for table in list(tables):
        table._element.getparent().remove(table._element)
    for paragraph in list(paragraphs):
        paragraph._element.getparent().remove(paragraph._element)


def set_cell_borders(cell, *, bottom: bool = False, color: str = "C9D1D9", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "right", "insideH", "insideV"):
        el = tc_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        el.set(qn("w:val"), "nil")

    bottom_el = tc_borders.find(qn("w:bottom"))
    if bottom_el is None:
        bottom_el = OxmlElement("w:bottom")
        tc_borders.append(bottom_el)
    if bottom:
        bottom_el.set(qn("w:val"), "single")
        bottom_el.set(qn("w:sz"), size)
        bottom_el.set(qn("w:color"), color)
    else:
        bottom_el.set(qn("w:val"), "nil")


def set_cell_margins(cell, top=40, start=30, bottom=40, end=30) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_field(paragraph) -> None:
    run_prefix = paragraph.add_run("Стр. ")
    set_run_font(run_prefix, size=9.5, color=RGBColor(90, 90, 90))

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    page_text = OxmlElement("w:t")
    page_text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    page_run = paragraph.add_run()
    set_run_font(page_run, size=9.5, color=RGBColor(90, 90, 90))
    page_run._r.append(begin)
    page_run._r.append(instr)
    page_run._r.append(separate)
    page_run._r.append(page_text)
    page_run._r.append(end)

    run_middle = paragraph.add_run(" из ")
    set_run_font(run_middle, size=9.5, color=RGBColor(90, 90, 90))

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " NUMPAGES "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    pages_text = OxmlElement("w:t")
    pages_text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    pages_run = paragraph.add_run()
    set_run_font(pages_run, size=9.5, color=RGBColor(90, 90, 90))
    pages_run._r.append(begin)
    pages_run._r.append(instr)
    pages_run._r.append(separate)
    pages_run._r.append(pages_text)
    pages_run._r.append(end)


def set_table_widths(table, widths_inches: list[float]) -> None:
    widths_dxa = [int(width * 1440) for width in widths_inches]
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = tbl.tblGrid
    if tbl_grid is not None:
        tbl.remove(tbl_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    tbl.insert(1, tbl_grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))


def build_header(section) -> None:
    header = section.header
    clear_container(header.paragraphs, header.tables)

    table = header.add_table(rows=1, cols=2, width=Inches(6.9))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_widths(table, [3.9, 3.0])

    left, right = table.rows[0].cells
    for cell in (left, right):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=30, bottom=55)
        set_cell_borders(cell, bottom=True)

    p_left = left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_left.add_run()
    run.add_picture(str(LOGO_PATH), width=Inches(2.0))

    p_right = right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_right.add_run("Партнерское предложение")
    set_run_font(run, size=10, color=RGBColor(90, 90, 90), bold=False)


def build_footer(section) -> None:
    footer = section.footer
    clear_container(footer.paragraphs, footer.tables)

    table = footer.add_table(rows=1, cols=3, width=Inches(6.9))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_widths(table, [1.5, 3.9, 1.5])

    left, center, right = table.rows[0].cells
    for cell in (left, center, right):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=50, bottom=20)
        set_cell_borders(cell, bottom=False)

    p_left = left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_left.add_run("Светон")
    set_run_font(run, size=9.5, color=RGBColor(90, 90, 90))

    p_center = center.paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_center.add_run("Коммерческое предложение для партнеров")
    set_run_font(run, size=9.5, color=RGBColor(90, 90, 90))

    p_right = right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(p_right)


def main() -> None:
    doc = Document(str(DOCX_PATH))
    for section in doc.sections:
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.35)
        build_header(section)
        build_footer(section)
    doc.save(str(DOCX_PATH))


if __name__ == "__main__":
    main()
