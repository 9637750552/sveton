from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "01_docs/legal/Договор_о_привлечении_и_сопровождении_клиентов_Светон.md"
DOCX_PATH = ROOT / "01_docs/legal/Договор_о_привлечении_и_сопровождении_клиентов_Светон.docx"


def set_run_font(run, size: float = 10.5, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), "Times New Roman")


def set_paragraph(paragraph, *, before=0, after=4, line=1.05, align=None) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_text_paragraph(doc: Document, text: str, *, size=10.5, bold=False, italic=False, align=None, after=4):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, after=after, align=align)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return paragraph


def clear_cell(cell) -> None:
    for paragraph in list(cell.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, color="FFFFFF") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")
        el.set(qn("w:color"), color)


def set_table_widths(table, widths_inches: list[float]) -> None:
    widths_dxa = [int(width * 1440) for width in widths_inches]
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))


def parse_requisites(lines: list[str], start_idx: int):
    customer: list[str] = []
    contractor: list[str] = []
    signatures: list[str] = []
    current = None
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if line == "Заказчик:":
            current = customer
        elif line == "Исполнитель:":
            current = contractor
        elif line == "Подписи:":
            current = signatures
        elif line:
            if current is not None:
                current.append(line.rstrip("  "))
        i += 1
    return customer, contractor, signatures


def add_requisites_table(doc: Document, customer: list[str], contractor: list[str], signatures: list[str]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_widths(table, [3.25, 3.25])

    for cell, title, rows in (
        (table.rows[0].cells[0], "Заказчик", customer),
        (table.rows[0].cells[1], "Исполнитель", contractor),
    ):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_margins(cell, top=80, start=90, bottom=80, end=90)
        set_cell_borders(cell)
        clear_cell(cell)
        p = cell.add_paragraph()
        set_paragraph(p, after=4)
        run = p.add_run(title)
        set_run_font(run, bold=True)
        for row in rows:
            p = cell.add_paragraph()
            set_paragraph(p, after=2, line=1.0)
            run = p.add_run(row)
            set_run_font(run, size=10)

    doc.add_paragraph()
    if signatures:
        add_text_paragraph(doc, "Подписи:", bold=True, after=6)
        for line in signatures:
            add_text_paragraph(doc, line, after=8)


def build_docx() -> None:
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.5)

    in_requisites = False
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue

        if line.startswith("# "):
            add_text_paragraph(doc, line[2:], size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
        elif line.startswith("## "):
            heading = line[3:]
            add_text_paragraph(doc, heading, size=11.5, bold=True, after=6)
            if heading == "8. Реквизиты и подписи Сторон":
                customer, contractor, signatures = parse_requisites(lines, i + 1)
                add_requisites_table(doc, customer, contractor, signatures)
                break
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style=None)
            set_paragraph(paragraph, after=2, line=1.0)
            paragraph.paragraph_format.left_indent = Cm(0.55)
            paragraph.paragraph_format.first_line_indent = Cm(-0.25)
            run = paragraph.add_run("- " + line[2:])
            set_run_font(run)
        else:
            text = line.replace("  ", "")
            is_clause = re.match(r"^\d+\.\d+\.", text) is not None
            paragraph = add_text_paragraph(doc, text, bold=False, after=4 if is_clause else 5)
            if is_clause:
                paragraph.paragraph_format.first_line_indent = Cm(0.0)
        i += 1

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
