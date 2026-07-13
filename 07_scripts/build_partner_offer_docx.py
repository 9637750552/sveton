from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "01_docs/legal/Коммерческое_предложение_партнерам_Светон.md"
DOCX_PATH = ROOT / "01_docs/legal/Коммерческое_предложение_партнерам_Светон.docx"
SVETON_LOGO = ROOT / "01_docs/legal/assets/Logo_Sveton_IBP_2_white.png"
BINEOS_LOGO = ROOT / "01_docs/legal/assets/BINEOS_logo.png"


HYPERLINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
IMG_RE = re.compile(r'<img\s+[^>]*src="([^"]+)"[^>]*alt="([^"]*)"[^>]*height="([^"]+)"[^>]*>')


def parse_markdown() -> dict:
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    i = 0

    while i < len(lines) and not lines[i].startswith("# "):
        i += 1

    title = lines[i][2:].strip()
    i += 1
    while i < len(lines) and not lines[i].strip():
        i += 1
    subtitle = lines[i].strip()
    i += 1
    while i < len(lines) and not lines[i].strip():
        i += 1
    intro = lines[i].strip()
    i += 1

    sections: list[dict] = []
    while i < len(lines):
        if not lines[i].startswith("## "):
            i += 1
            continue
        heading = lines[i][3:].strip()
        i += 1
        while i < len(lines) and not lines[i].strip():
            i += 1

        paragraphs: list[str] = []
        table: list[list[str]] | None = None
        list_items: list[str] = []

        block: list[str] = []
        while i < len(lines) and not lines[i].startswith("## "):
            line = lines[i]
            if line.startswith("|"):
                if block:
                    paragraphs.append(" ".join(x.strip() for x in block))
                    block = []
                raw_rows: list[str] = []
                while i < len(lines) and lines[i].startswith("|"):
                    raw_rows.append(lines[i])
                    i += 1
                table = []
                for row in raw_rows:
                    cells = [cell.strip() for cell in row.strip().strip("|").split("|")]
                    if all(set(cell) <= {"-"} for cell in cells):
                        continue
                    table.append(cells)
                continue
            if re.match(r"^\d+\.\s", line):
                if block:
                    paragraphs.append(" ".join(x.strip() for x in block))
                    block = []
                while i < len(lines) and re.match(r"^\d+\.\s", lines[i]):
                    list_items.append(re.sub(r"^\d+\.\s*", "", lines[i]).strip())
                    i += 1
                continue
            if not line.strip():
                if block:
                    paragraphs.append(" ".join(x.strip() for x in block))
                    block = []
                i += 1
                continue
            block.append(line.strip())
            i += 1
        if block:
            paragraphs.append(" ".join(x.strip() for x in block))

        sections.append(
            {
                "heading": heading,
                "paragraphs": paragraphs,
                "table": table,
                "list_items": list_items,
            }
        )

    return {
        "title": title,
        "subtitle": subtitle,
        "intro": intro,
        "sections": sections,
    }


def set_cell_margins(cell, top=60, start=50, bottom=60, end=50) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_borders(cell, top=False, bottom=False) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("left", "right", "insideH", "insideV"):
        el = tc_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        el.set(qn("w:val"), "nil")

    for edge, enabled in (("top", top), ("bottom", bottom)):
        el = tc_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        if enabled:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "8")
            el.set(qn("w:color"), "AAB7C4")
        else:
            el.set(qn("w:val"), "nil")


def set_table_layout_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_table_widths(table, widths_inches: list[float]) -> None:
    widths_dxa = [int(width * 1440) for width in widths_inches]

    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_grid = tbl.tblGrid
    if tbl_grid is not None:
        tbl.remove(tbl_grid)
    tbl_grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    tbl.insert(1, tbl_grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_inches[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))


def set_row_no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def style_runs(paragraph, size=12.5, color=RGBColor(0, 0, 0), bold=False) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.08
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color


def add_hyperlink(paragraph, text, url) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_inline_markup(paragraph, text: str) -> None:
    pos = 0
    while pos < len(text):
        img_match = IMG_RE.search(text, pos)
        link_match = HYPERLINK_RE.search(text, pos)

        matches = [m for m in (img_match, link_match) if m]
        if not matches:
            paragraph.add_run(text[pos:])
            break

        match = min(matches, key=lambda m: m.start())
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])

        if match.re is IMG_RE:
            src = match.group(1)
            logo_path = MD_PATH.parent / src
            run = paragraph.add_run()
            run.add_picture(str(logo_path), height=Pt(11))
        else:
            add_hyperlink(paragraph, match.group(1), match.group(2))
        pos = match.end()


def add_paragraph(doc: Document, text: str, size=12.5, color=RGBColor(0, 0, 0), after=6) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.08
    add_inline_markup(paragraph, text)
    style_runs(paragraph, size=size, color=color)


def build_docx() -> None:
    data = parse_markdown()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal_style.font.size = Pt(12)

    title_blue = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    shape = p.add_run().add_picture(str(SVETON_LOGO), width=Inches(5.8))
    shape._inline.docPr.set("name", "")
    shape._inline.docPr.set("descr", "")
    shape._inline.graphic.graphicData.pic.nvPicPr.cNvPr.set("name", "")
    shape._inline.graphic.graphicData.pic.nvPicPr.cNvPr.set("descr", "")

    add_paragraph(doc, data["title"], size=26, color=title_blue, after=10)
    add_paragraph(doc, data["subtitle"], size=13, after=9)
    add_paragraph(doc, data["intro"], size=12.5, after=12)

    for section_data in data["sections"]:
        add_paragraph(doc, section_data["heading"], size=22, color=title_blue, after=8)
        for paragraph_text in section_data["paragraphs"]:
            add_paragraph(doc, paragraph_text, size=12.5, after=7)

        table = section_data["table"]
        if table:
            if section_data["heading"] == "Модель работы Светон":
                widths = [1.25, 5.55]
            elif section_data["heading"] == "Роли партнера в сделке":
                widths = [1.2, 3.35, 2.25]
            else:
                widths = [0.9, 3.65, 2.25]

            is_model = section_data["heading"] == "Модель работы Светон"
            row_count = len(table) if is_model else len(table)
            full_table = doc.add_table(rows=row_count, cols=len(widths))
            full_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            full_table.autofit = False
            set_table_layout_fixed(full_table)
            set_table_widths(full_table, widths)

            for row_idx, row_data in enumerate(table):
                row = full_table.rows[row_idx]
                set_row_no_split(row)
                is_header = not is_model and row_idx == 0
                for idx, text in enumerate(row_data):
                    cell = row.cells[idx]
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    set_cell_margins(cell, top=55 if is_header else 50, bottom=55 if is_header else 50)
                    set_borders(cell, top=is_header, bottom=True)
                    paragraph = cell.paragraphs[0]
                    add_inline_markup(paragraph, text)
                    style_runs(paragraph, size=12.5, bold=is_header)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        if section_data["list_items"]:
            for index, item in enumerate(section_data["list_items"], start=1):
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.35)
                paragraph.paragraph_format.first_line_indent = Inches(-0.15)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.08
                paragraph.add_run(f"{index}. ")
                add_inline_markup(paragraph, item)
                style_runs(paragraph, size=12.5)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
