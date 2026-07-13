from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
LOGO = ROOT / "07_forms" / "assets" / "sveton-logo.png"
OUT = ROOT / "02_output" / "legal" / "Коммерческое_предложение_партнерам_Светон.docx"

BLUE = "1F4D78"
BLUE_ACCENT = "2E74B5"
INK = "1F1F1F"
MUTED = "666666"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "D7DBE2"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_inches: list[float]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths_inches) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_text(cell, text: str, *, bold=False, color=INK, size=10.2) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def set_column_widths(table, widths_inches: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_column_widths(table, widths)

    header = table.rows[0]
    set_repeat_table_header(header)
    set_row_cant_split(header)
    for i, text in enumerate(headers):
        set_cell_shading(header.cells[i], LIGHT_BLUE)
        set_cell_text(header.cells[i], text, bold=True, color=BLUE, size=10)

    for row_data in rows:
        row = table.add_row()
        set_row_cant_split(row)
        for i, text in enumerate(row_data):
            if i == 0:
                set_cell_shading(row.cells[i], LIGHT_GRAY)
                set_cell_text(row.cells[i], text, bold=True, color=BLUE, size=10)
            else:
                set_cell_text(row.cells[i], text, size=10.2)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 2"
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(BLUE_ACCENT)
    run.bold = True


def add_numbered_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        p.add_run(item)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        p.add_run(item)


def set_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Светон | партнерское предложение")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    set_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    if LOGO.exists():
        p_logo = doc.add_paragraph()
        p_logo.paragraph_format.space_after = Pt(14)
        p_logo.add_run().add_picture(str(LOGO), width=Inches(2.45))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Партнерское предложение Светон")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Партнерская программа по привлечению и сопровождению клиентов")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Версия для обсуждения коммерческих условий после заполнения анкеты партнера")
    run.font.size = Pt(9.5)
    run.italic = True
    run.font.color.rgb = RGBColor.from_string(MUTED)

    add_h2(doc, "Смысл сотрудничества")
    add_table(
        doc,
        ["Вопрос", "Ответ"],
        [
            ["Для кого", "Для владельцев домов, коммерческих объектов и небольших предприятий, где электропитание связано с комфортом, безопасностью или непрерывностью работы."],
            ["Зачем", "Чтобы у клиента был управляемый запас автономности на случай отключений, аварий и нестабильной сети."],
            ["Что делаем", "Подбираем и запускаем системы на базе инверторов, аккумуляторов, ИБП, стабилизаторов и сопутствующего оборудования."],
            ["Как делаем", "Сначала разбираем задачу и объект, затем подбираем решение, поставляем оборудование, организуем монтаж и ввод в эксплуатацию."],
            ["В чем уникальность", "17 лет на рынке, инженерная экспертиза, практический опыт объектов и ответственность за работающий результат."],
            ["С кем", "С партнерами рядом с клиентом: электриками, монтажными бригадами, инженерами, строителями, проектировщиками, прорабами и компаниями."],
        ],
        [1.55, 4.95],
    )

    add_h2(doc, "Роли партнера в сделке")
    add_table(
        doc,
        ["Роль", "Что делает партнер", "Кто оплачивает"],
        [
            ["Привлечение клиента", "Передает клиента, у которого есть потребность в резервном или автономном электропитании, и помогает Светон установить первый контакт.", "Светон выплачивает вознаграждение за привлечение клиента."],
            ["Технический осмотр и сбор исходных данных", "Проверяет условия подключения, щиты, ввод, место установки, нагрузки, ограничения, делает фото/видео и передает исходные данные для расчета и подбора решения.", "Светон выплачивает вознаграждение, если осмотр согласован со Светон."],
            ["Монтажные и иные работы", "Выполняет подготовительные работы, монтаж, подключение оборудования, участвует в запуске системы.", "Клиент оплачивает работы напрямую партнеру."],
        ],
        [1.45, 3.15, 1.9],
    )

    add_h2(doc, "Форматы партнерства по договору")
    add_body_paragraph(doc, "Для договора о привлечении и сопровождении клиентов используются два формата участия:")
    add_table(
        doc,
        ["Формат", "Что входит", "Кто платит"],
        [
            ["Тип A", "Партнер привлекает клиента и передает его контакты.", "Светон выплачивает вознаграждение за привлечение клиента."],
            ["Тип B", "Партнер привлекает клиента, самостоятельно проводит технический осмотр и передает исходные данные для расчета и подбора решения.", "Светон выплачивает вознаграждение за привлечение клиента и согласованный технический осмотр."],
        ],
        [0.85, 3.85, 1.8],
    )
    add_body_paragraph(
        doc,
        "Если по объекту возникает необходимость монтажа, подготовительных или иных работ, такие работы согласуются между клиентом и партнером отдельно и оплачиваются клиентом напрямую партнеру.",
    )

    add_h2(doc, "Как фиксируется участие партнера")
    add_body_paragraph(
        doc,
        "Участие партнера фиксируется по конкретной сделке. Мы не считаем партнера закрепленным за клиентом вообще и не начисляем вознаграждение за клиентов, которые уже были в базе Светон или находились в работе до участия партнера.",
    )
    add_body_paragraph(
        doc,
        "Формат участия, клиент и сделка фиксируются в CRM и/или в реестре сделок. Это нужно, чтобы стороны одинаково понимали, за какую сделку и за какой формат участия возникает вознаграждение.",
    )

    add_h2(doc, "На каких условиях мы работаем")
    add_body_paragraph(
        doc,
        "Вознаграждение за привлечение клиента и согласованный технический осмотр выплачивает Светон после того, как система доведена до результата и введена в эксплуатацию.",
    )
    add_body_paragraph(
        doc,
        "Размер вознаграждения рассчитывается по номинальной мощности инвертора. Мощность аккумуляторных батарей в расчет не включается. Если мощность инвертора дробная, расчет выполняется пропорционально фактической номинальной мощности.",
    )
    add_body_paragraph(
        doc,
        "Выплаты производятся только безналично по реквизитам партнера. Конкретные ставки, порядок расчета и юридические условия согласовываются после заполнения анкеты партнера и указываются в договоре при подписании.",
    )

    add_h2(doc, "Чего партнер не делает")
    add_body_paragraph(
        doc,
        "Партнер не является представителем Светон, не подписывает документы от имени Светон, не принимает обязательства перед клиентом за Светон и не согласует коммерческие условия от нашего имени.",
    )
    add_body_paragraph(
        doc,
        "Светон самостоятельно ведет переговоры с клиентом, готовит коммерческое предложение, заключает договор, принимает оплату, поставляет оборудование и отвечает за согласованную с клиентом часть результата.",
    )

    add_h2(doc, "Как начать сотрудничество")
    add_numbered_list(
        doc,
        [
            "Ознакомиться с партнерским предложением.",
            "Заполнить анкету партнера.",
            "Передать реквизиты для договора и выплат.",
            "Согласовать формат участия и условия.",
            "Подписать договор о привлечении и сопровождении клиентов.",
            "Передавать клиентов или участвовать в сделках по согласованному порядку.",
        ],
    )

    add_h2(doc, "Документы и ссылки")
    add_body_paragraph(doc, "К партнерскому предложению прикладываются:")
    add_bullets(
        doc,
        [
            "договор о привлечении и сопровождении клиентов;",
            "анкета партнера;",
            "ссылка на сайт Светон.",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
